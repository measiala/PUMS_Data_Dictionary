#!/usr/bin/python3

import argparse
import os

import logging

from docx import Document
from docx.shared import Inches, Pt

from class_defs import *
from process_arguments import *
from classify_line import classify_line
from process_line import process_line

from docx_output import initialize_doc
from output_all import *

INDIR = './input'
OUTDIR = './output'
    
argp = argparse.ArgumentParser(description='Provide input data dictionary parameters')

argp.add_argument('yyyy',metavar='YYYY',type=int,
                  help='Ending 4-digit data year for data dictionary')
argp.add_argument('period',metavar='PERIOD',type=int,choices=[1,3,5],
                  help='Length of period (1/3/5)')
argp.add_argument('--ifile','-i',metavar='InFile',type=str,
                  help='Path to input data dictionary')
argp.add_argument('--outdir','-o',metavar='OutDir',type=str,
                  help='Path to output directory')
argp.add_argument('--replace','-r',action='store_true',
                  help='Replace existing output files')
argp.add_argument('--loglvl','-l',metavar='LogLevel',type=str,
                  choices=['ERROR','WARNING','INFO'],
                  help='Log Output Level (ERROR/WARNING/INFO)')

args = argp.parse_args()

yyyy = assign_yyyy(args)
if yyyy == None:
    exit(2)
    
period = assign_period(args)
if period == None:
    exit(2)

out = assign_infile(args,INDIR)
if out == None:
    exit(2)
else:
    [ infile, basefile ] = out

out = assign_outfiles(args,OUTDIR)
if out == None:
    exit(2)
else:
    [ outdoc, outtxt, outcsv, outlog ] = out

""" Setup logging """    
print("NOTE: Writing output LOG to: %s" % outlog)

if args.loglvl:
    num_lvl = getattr(logging, args.loglvl.upper(), None)
else:
    num_lvl = getattr(logging, 'INFO', None)
logging.basicConfig(filename=outlog,filemode = 'w',level=num_lvl)

""" Print to stdout basic information """
logging.info("-----")
logging.info("Reading input DOCX from: %s" % infile)
logging.info("Reading input PageFooter.docx from: %s" % basefile)
logging.info("-----")
logging.info("Writing output DOCX to: %s" % outdoc)
logging.info("Writing output TXT to: %s" % outtxt)
logging.info("Writing output CSV to: %s" % outcsv)
logging.info("-----")

""" Setup base output Word document """
ddword = Document(basefile)
initialize_doc(ddword)
ddword.core_properties.author = 'US Census Bureau'
ddword.core_properties.title = os.path.splitext(os.path.basename(outdoc))[0]

""" Open text and csv files for writing """
ddtext = open(outtxt,'w')
ddcsv = open(outcsv,'w')

""" Begin read of input Word document """
ddorig = Document(infile)
origpars = ddorig.paragraphs

""" Initialize Dictionary Storage """
dd = DataDict('PUMS Variable Data Dictionary')
pl = PUMSDict('PUMS Layouts by Record Type')

""" Initialize lag variables """
rt = ''
srt = ''
pltype = 'Blank'
pvar = ''
pval = ''

""" Iterate over all paragraphs in document """
for par in origpars:
    """
    The loop first classifies the line based on the current line and the preceding line type.
    Next, we process the line once we know what type of data it is.
    Lastly we initialize the previous line type variable for the next pass.
    """
    import re
    lfparser = re.compile(r"[\t ]*\n[\t .]*")
    
    p = par.text.replace(chr(160),' ').\
        replace(chr(8212),'-').\
        replace(chr(8220),'"').\
        replace(chr(8221),'"').\
        replace(chr(8216),"'").\
        replace(chr(8217),"'").strip()

    p = ' '.join(lfparser.split(p))

    ltype = classify_line(p,pltype)
    if ltype != None:
        tmp = process_line(p,ltype,dd,pl,rt,srt,pvar,pval)
        if tmp != None:
            if ltype == 'Header':
                if len(tmp) == 1:
                    rt = tmp
                    srt = ''
                elif len(tmp) > 1:
                    srt = tmp
            elif ltype == 'Var Name':
                pvar = tmp
            elif ltype == 'Var Value':
                pval = tmp
            pltype = ltype
        else:
            logging.error("Unable to process line '%s'. Please fix.",p)
            logging.error("Last variable was %s.",pvar)
            exit(2)
    else:
        logging.error("Line '%s' is not resolvable. Please fix.",p)
        logging.error("Last variable was %s.",pvar)
        exit(2)
        
output_var_block(pl,dd,ddtext,ddword,ddcsv,custom='YES')

ddtext.close()
ddcsv.close()

ddword.save(outdoc)
