#!/usr/bin/python3

from docx           import Document
from docx.shared    import Inches, Pt
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER, WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

import textwrap

from hyperlink import add_hyperlink

#############################
### Setup output document ###
#############################

def initialize_doc( doc ):
    ### Set 1-inch page margins
    page_width = 6.5 * 72.27 / 72
    margin = (8.5 - page_width) / 2
    doc.sections.left_margin  = Inches(margin)
    doc.sections.right_margin = Inches(margin)

    ### Define styles
    normal = doc.styles['Normal']
    cnorm  = doc.styles.add_style('Normal Center', WD_STYLE_TYPE.PARAGRAPH)
    head1  = doc.styles['Heading 1']
    head2  = doc.styles['Heading 2']
    head3  = doc.styles['Heading 3']
    vardes = doc.styles.add_style('Var Desc',  WD_STYLE_TYPE.PARAGRAPH)
    varval = doc.styles.add_style('Var Value', WD_STYLE_TYPE.PARAGRAPH)
    note   = doc.styles.add_style('Note',      WD_STYLE_TYPE.PARAGRAPH)

    def style_props( style, priority, align, ptindent, ptbefore, ptafter, bold, color = 'None' ):
        if style != normal:
            style.base_style = normal
        style.hidden = False
        style.quick_style = True
        style.priority = priority
        style.font.name = 'Courier New'
        style.font.size = Pt(10)
        style.font.bold = bold
        if align == 'CENTER':
            style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            style.paragraph_format.left_indent = Pt(ptindent)
        style.paragraph_format.space_before = Pt(ptbefore)
        style.paragraph_format.space_after  = Pt(ptafter)
        style.paragraph_format.widow_control = True
        if 1 == 0: #color != 'None':
            style.font.color = "5B9BC5"
        return;
    
    # Normal Style
    style_props( normal, 0, 'LEFT', 0, 0, 0, False)

    # Normal Center Style
    style_props( cnorm, 7, 'CENTER', 0, 0, 0, False )
    
    # Heading 1 Style for title
    style_props( head1, 1, 'CENTER', 0, 24, 0, True, '5B9BC5' )

    # Heading 2 Style for Record Type Seperators
    style_props( head2, 2, 'CENTER', 0, 12, 12, True, '5B9BC5' )

    # Heading 3 Style for Variable header
    style_props( head3, 3, 'LEFT', 0, 12, 0, True, '5B9BC5' )
    head3.paragraph_format.tab_stops.add_tab_stop(Inches(1.0))

    # Variable Description Style
    style_props( vardes, 4, 'LEFT', 36, 0, 0, False )

    # Legal Value Style
    style_props( varval, 5, 'LEFT', 72, 0, 0, False )

    # Note
    style_props( note, 6, 'LEFT', 0, 12, 0, False )
    return;

def docx_out_title( title, doc ):
    doc.add_heading( title, level=1 )
    return;

def docx_out_date( reldate, doc ):
    doc.add_paragraph( reldate, style='Normal Center' )
    return;

def docx_out_major_record_type( rtype, doc ):
    doc.add_heading( rtype, level=1 )
    return;

def docx_out_minor_record_type( rtype, doc ):
    doc.add_heading( rtype, level=2 )
    return

def docx_out_var_name( var_name, var_len, doc ):
    doc.add_heading( var_name + '\t' + str(var_len), level=3 )
    return;

def docx_out_var_desc( var_desc, doc ):
    doc.add_paragraph( var_desc, style='Var Desc' )
    return;

def docx_out_var_val(var_val, var_val_desc, var_len, doc, wrap_text = 'YES' ):
    if wrap_text == 'YES':
        tabchar = 12 + 2*int(var_len) + 3
        tabloc = tabchar*6
        max_wd = 78 - tabchar - 2
        lines = textwrap.wrap(textwrap.dedent(var_val_desc).strip(), width=max_wd)
        # if var_val == 'bbbbb':
        #     print(tabchar, tabloc, max_wd, len(lines), len(lines[0]), var_val_desc )
        if len(lines) == 1:
            newp = doc.add_paragraph( var_val + '\t.' + lines[0], style='Var Value' )
            newp.paragraph_format.tab_stops.add_tab_stop( Pt(tabloc) )
        elif len(lines) > 1:
            oline = var_val + '\t.' + lines[0] + '\n'
            for i in range(1,len(lines)-1):
                oline = oline + '\t.' + lines[i] + '\n'
            oline = oline + '\t.' + lines[-1]
            newp = doc.add_paragraph( oline, style = 'Var Value' )
            newp.paragraph_format.tab_stops.add_tab_stop( Pt(tabloc) )
    else:
        newp = doc.add_paragraph( var_val + '\t.' + var_val_desc, style='Var Value' )
    # Each char has length 6pt so below allows two values of var_len plus 1 inch plus the two dots
    newp.paragraph_format.tab_stops.add_tab_stop( Pt(6*(2*int(var_len) + 3) + 72) )
    return;

def docx_out_note( note, doc, urls = 'NO' ):
    if urls == 'YES':
        outnote = ''
        outp = doc.add_paragraph('', style='Note')
        len_note = len(note)
        note_pos = 0

        while note_pos < len_note:
            if note.find("http",note_pos) == -1:
                outnote = note[note_pos:]
                outp.add_run(outnote)
                note_pos = len_note
            else:
                old_note_pos = note_pos
                note_pos = note.find("http", old_note_pos)
                outnote = note[old_note_pos:note_pos]
                link_start = note[note_pos:]
                outp.add_run(outnote)
                if len(link_start.split()) == 1:
                    if link_start[-1] == '.':
                        link = link_start[:(len(link_start) - 1)]
                    else:
                        link = link_start
                else:
                    link = link_start.split()[0]
                note_pos = note_pos + len(link)
                hyperlink = add_hyperlink(outp, link, link, '0000FF', True)
    else:
        doc.add_paragraph( note, style='Note' )
    return;
